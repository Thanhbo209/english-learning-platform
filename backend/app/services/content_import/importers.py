import csv
import io
from dataclasses import dataclass, field

import openpyxl
import pypdf
from docx import Document


@dataclass
class RawContent:
    """Format-generic extraction result. Importers never know about content types."""

    text: str | None = None
    rows: list[dict[str, str]] | None = None
    tables: list[list[list[str]]] = field(default_factory=list)
    sheets: dict[str, list[dict[str, str]]] = field(default_factory=dict)


class FileImportError(Exception):
    """Raised when a file cannot be read as its declared format."""


def import_docx(data: bytes) -> RawContent:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise FileImportError("Could not read this file as a .docx document.") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
    ]
    return RawContent(text="\n".join(paragraphs), tables=tables)


def import_xlsx(data: bytes, sheet_name: str | None = None) -> RawContent:
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    except Exception as exc:
        raise FileImportError("Could not read this file as an .xlsx spreadsheet.") from exc

    sheets: dict[str, list[dict[str, str]]] = {}
    is_multi_sheet = len(workbook.worksheets) > 1

    for sheet in workbook.worksheets:
        title = sheet.title or "Sheet"
        rows_iter = sheet.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            sheets[title] = []
            continue

        header = [str(cell).strip() if cell is not None else "" for cell in header_row]
        sheet_rows = []
        for raw_row in rows_iter:
            if all(cell is None for cell in raw_row):
                continue
            row = {}
            for column_name, cell in zip(header, raw_row, strict=False):
                if column_name:
                    row[column_name] = "" if cell is None else str(cell)
            if row:
                if is_multi_sheet or sheet_name == "all":
                    row["_sheet_name"] = title
                sheet_rows.append(row)
        sheets[title] = sheet_rows

    if not sheets:
        return RawContent(rows=[], sheets={})

    if sheet_name and sheet_name in sheets:
        selected_rows = sheets[sheet_name]
    elif sheet_name == "all":
        selected_rows = []
        for s_rows in sheets.values():
            selected_rows.extend(s_rows)
    else:
        non_empty = [s_rows for s_rows in sheets.values() if s_rows]
        if len(sheets) == 1:
            selected_rows = list(sheets.values())[0]
        elif non_empty:
            selected_rows = []
            for s_rows in non_empty:
                selected_rows.extend(s_rows)
        else:
            selected_rows = []

    return RawContent(rows=selected_rows, sheets=sheets)


def import_pdf(data: bytes) -> RawContent:
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise FileImportError("Could not read this file as a .pdf document.") from exc
    return RawContent(text=text)


def import_csv(data: bytes) -> RawContent:
    text: str | None = None
    encodings = ("utf-8-sig", "utf-8", "cp1252", "latin-1")
    for encoding in encodings:
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if text is None:
        raise FileImportError(
            "Could not read this file as text. Please ensure it is UTF-8 or Windows-1252 encoded."
        )

    try:
        reader = csv.DictReader(io.StringIO(text))
        rows = [{(k or "").strip(): (v or "").strip() for k, v in row.items()} for row in reader]
    except Exception as exc:
        raise FileImportError("Could not parse CSV content.") from exc

    return RawContent(rows=rows)


IMPORTERS = {
    "docx": import_docx,
    "xlsx": import_xlsx,
    "pdf": import_pdf,
    "csv": import_csv,
}


def import_file(source_format: str, data: bytes, sheet_name: str | None = None) -> RawContent:
    importer = IMPORTERS.get(source_format)
    if importer is None:
        raise FileImportError(f"Unsupported file format: {source_format}")
    if source_format == "xlsx":
        return importer(data, sheet_name=sheet_name)
    return importer(data)
